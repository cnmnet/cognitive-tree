#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""对指定报告文件（Markdown / Word）执行质量分与认知层级分双轨评分。"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path
from typing import Any, Dict

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from core.scoring import (
    score_report,
    score_summary,
    score_surprise_winning,
)
from external.services import score_cognitive_level


DEFAULT_OUT_DIR = ROOT / "docs" / "输出" / "评分"


def read_report(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8")


def build_quality(path: Path, text: str, source: str) -> Dict[str, Any]:
    scores = score_report(text)
    return {
        "method": "local_heuristic",
        "dimensions": scores,
        "surprise_winning": score_surprise_winning(text),
        "weighted_total": score_summary(scores),
        "reference_only_for_external": source == "external_report",
    }


def build_cognitive(text: str, enabled: bool) -> Dict[str, Any]:
    if not enabled:
        return {
            "method": "llm_assisted",
            "status": "not_requested",
        }
    return score_cognitive_level(text)


def to_markdown(rows: list, tag: str) -> str:
    headers = [
        "报告",
        "论证深度",
        "证据质量",
        "逻辑严谨",
        "视角多样",
        "创新洞察",
        "出奇制胜",
        "结构组织",
        "可读性",
        "实用价值",
        "加权总分",
    ]
    lines = [
        f"# {tag}",
        "",
        "- 评分系统：core/scoring.py 本地九维启发式评分（零依赖）",
        "- 说明：仅作内部趋势参考；外部 Word 报告质量分只作参考，不参与认知层级判定",
        "",
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        s = row["scores"]
        values = [
            row["file"],
            s["argument_depth"],
            s["evidence_quality"],
            s["logic_rigor"],
            s["perspective_diversity"],
            s["innovation_insight"],
            s["surprise_winning"],
            s["structure_organization"],
            s["readability"],
            s["practical_value"],
            row["total"],
        ]
        lines.append("| " + " | ".join(str(v) for v in values) + " |")
    lines.append("")
    return "\n".join(lines)


def cognitive_markdown(rows: list, tag: str) -> str:
    lines = [
        f"# {tag} · 认知层级",
        "",
        "- 评分系统：external/services.py 大模型七维认知评分",
        "",
        "| 报告 | 状态 | 加权总分 | 认知层级 | 出奇制胜标签 |",
        "| --- | --- | ---: | --- | --- |",
    ]
    for row in rows:
        cog = row["cognitive"]
        status = cog.get("status", "-")
        total = cog.get("weighted_total")
        total = total if total is not None else "-"
        level = cog.get("cognitive_level")
        level = level if level is not None else "-"
        tags = "、".join(cog.get("strategy_tags", [])) or "-"
        lines.append(
            f"| {row['file']} | {status} | {total} | {level} | {tags} |"
        )
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="质量分与认知层级分双轨评分汇总")
    parser.add_argument("--file", type=Path, action="append", required=True, help="报告文件（可多次）")
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR, help="输出目录")
    parser.add_argument("--tag", default="报告八维评分", help="输出文件标签")
    parser.add_argument("--cognitive", action="store_true", help="调用大模型生成认知层级分")
    parser.add_argument(
        "--source",
        choices=["system_report", "external_report"],
        default=None,
        help="报告来源；缺省按文件后缀判断",
    )
    args = parser.parse_args()

    rows = []
    for path in args.file:
        text = read_report(path)
        source = args.source
        if source is None:
            source = (
                "external_report"
                if path.suffix.lower() == ".docx"
                else "system_report"
            )
        quality = build_quality(path, text, source)
        cognitive = build_cognitive(text, args.cognitive)
        rows.append(
            {
                "file": path.name,
                "chars": len(text),
                "scores": quality["dimensions"],
                "total": quality["weighted_total"],
                "source": source,
                "quality": quality,
                "cognitive": cognitive,
            }
        )
    rows.sort(key=lambda r: r["total"], reverse=True)

    out_dir = args.out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = date.today().strftime("%Y%m%d")
    md_path = out_dir / f"{args.tag}_{stamp}.md"
    json_path = out_dir / f"{args.tag}_{stamp}.json"
    md_path.write_text(
        to_markdown(rows, args.tag) + "\n" + cognitive_markdown(rows, args.tag),
        encoding="utf-8",
    )
    json_path.write_text(
        json.dumps(
            {
                "schema": "crystal_tree.score.v2",
                "generated_at": stamp,
                "rows": rows,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(
        json.dumps(
            {"reports": len(rows), "md": str(md_path), "json": str(json_path)},
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
